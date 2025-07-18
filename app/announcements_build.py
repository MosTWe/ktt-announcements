from docx import Document, shared
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from datetime import datetime,date,timedelta

def write_announcements(zmanim):

    parsha = zmanim['parsha']
    
    weekstart = zmanim['sunday']['english_date']
    weekstart = datetime.strptime(weekstart, '%B %d').date() + timedelta(days=1)  # Assuming week starts on Sunday
    weekend = weekstart + timedelta(days=4)
    weekstart = weekstart.strftime('%m/%d')
    weekend = weekend.strftime('%m/%d')


    document = Document()
    section = document.sections[0]
    print(section.top_margin)
    section.top_margin = section.top_margin//2
    section.bottom_margin = section.bottom_margin//2
    section.left_margin = section.left_margin//2
    section.right_margin = section.right_margin//2

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'Erev Shabbos, {zmanim['friday']['english_month']} {zmanim['friday']['english_day']}').underline = True
    fri_english_ordinality = p.add_run(zmanim['friday']['english_ordinality'])
    fri_english_ordinality.font.superscript = True
    fri_english_ordinality.underline = True
    p.add_run(f' ({zmanim['friday']['hebrew_day']}')
    fri_hebrew_ordinality = p.add_run(zmanim['friday']['hebrew_ordinality'])
    fri_hebrew_ordinality.font.superscript = True
    fri_hebrew_ordinality.underline = True
    p.add_run(f' {zmanim['friday']['hebrew_month']})\n').underline = True
    p.add_run(f'{zmanim['friday']['early_shabbos']} Plag Mincha/Kabbalas Shabbos\n').bold = True
    p.add_run(f'Early Candle Lighting not before {zmanim['friday']['plag_mincha']}\n(The ideal time for Plag Candle Lighting is roughly 30 minutes after the start of Mincha)').italic = True
    p.add_run('\n')
    p.add_run(f'{zmanim['friday']['candle_lighting']} Candle Lighting/Mincha\n').bold = True
    p.add_run(f'{zmanim['friday']['shkia']} Shkia\n')
    p.add_run('\n')
    p.add_run(f'Shabbos, {zmanim['shabbos']['english_month']} {zmanim['shabbos']['english_day']}').underline = True
    shab_english_ordinality = p.add_run(zmanim['shabbos']['english_ordinality'])
    shab_english_ordinality.font.superscript = True
    shab_english_ordinality.underline = True
    p.add_run(f' ({zmanim['shabbos']['hebrew_day']}')
    shab_hebrew_ordinality = p.add_run(zmanim['shabbos']['hebrew_ordinality'])
    shab_hebrew_ordinality.font.superscript = True
    shab_hebrew_ordinality.underline = True
    p.add_run(f' {zmanim['shabbos']['hebrew_month']})\n').underline = True    
    p.add_run(f'{zmanim['shabbos']['shacharis']} Shacharis\n').bold = True
    p.add_run(f'{zmanim['shabbos']['zman_kriyas_shema']} Zman Kriyas Shema\n')
    p.add_run('\n')
    p.add_run('8:45 AM Youth Groups \n')
    p.add_run('9:30 AM Open Playroom for Toddlers \n')
    p.add_run('10:15 AM Mommy and Me (ages 4 and younger) \n')
    p.add_run('\n')
    p.add_run('Hot Kiddush\n').bold = True
    p.add_run('Thank you to our Kiddush Sponsors listed in the announcements!\n ')
    p.add_run('\n')
    p.add_run(f'{zmanim['shabbos']['mincha']} Mincha \n').bold = True
    p.add_run('Seudas Shelishis\n')
    p.add_run('Sponsored by _________ ').italic = True
    p.add_run('in honor of _________\n').italic = True
    p.add_run('Topic: _________\n').italic = True
    p.add_run('\n')
    p.add_run(f'{zmanim['shabbos']['shkia']} Shkia\n')
    p.add_run(f'{zmanim['shabbos']['eretz_yisrael_seder']} Eretz Yisrael Seder \n')
    p.add_run(f'{zmanim['shabbos']['havdalah']} Maariv/Havdalah\n').bold = True
    p.add_run('\n')
    p.add_run(f'Sunday, {zmanim['sunday']['english_month']} {zmanim['sunday']['english_day']}').underline = True
    sun_english_ordinality = p.add_run(zmanim['sunday']['english_ordinality'])
    sun_english_ordinality.font.superscript = True
    sun_english_ordinality.underline = True
    p.add_run(f' ({zmanim['sunday']['hebrew_day']}')
    sun_hebrew_ordinality = p.add_run(zmanim['sunday']['hebrew_ordinality'])
    sun_hebrew_ordinality.font.superscript = True
    sun_hebrew_ordinality.underline = True
    p.add_run(f' {zmanim['sunday']['hebrew_month']})\n').underline = True
    p.add_run(f'{zmanim['sunday']['shacharis']} Shacharis\n').bold = True
    p.add_run('\n')
    p.add_run(f'{zmanim['sunday']['zman_kriyas_shema']} Zman Kriyas Shema\n')
    p.add_run('Halacha Shiur & Breakfast\n')
    p.add_run('Sponsored by Steven Friedman\n').italic = True
    hebrew = p.add_run('לעילוי נשמת חיה רחל בת נחמן דוד שיבדל לחיים טובים\n')
    hebrew.italic = True
    hebrew.font.name = 'David'
    p.add_run('Topic: _________\n').italic = True
    p.add_run('\n')
    p.add_run(f'Weekday Shacharis ({weekstart} - {weekend})\n').underline = True
    p.add_run('6:30 AM ').bold = True
    p.add_run('Pre-Shacharis Shiur with Rabbi Davis\n')
    p.add_run('Coffee and light breakfast served daily, sponsored by Yael & Uri Segelman\n').italic = True
    p.add_run('6:35 AM ').bold = True 
    p.add_run('Monday and Thursday\n')
    p.add_run('6:45 AM ').bold = True
    last = p.add_run('Tuesday, Wednesday, and Friday\n')
    last.add_break(WD_BREAK.PAGE)
    heading = p.add_run('Announcements\n')
    heading.bold = True
    heading.underline = True
    heading.font.size = 16
    p.add_run('\n')
    p.add_run('Thank you to our Kiddush Sponsors this week:\n')
    p.add_run('\n')
    plat = p.add_run('Platinum Sponsorship\n')
    plat.bold = True
    plat.underline = True
    plat.font.color.rgb = shared.RGBColor(229, 228, 226) 
    p.add_run('\n')
    gold = p.add_run('Gold Sponsorship\n')
    gold.bold = True
    gold.underline = True
    gold.font.color.rgb = shared.RGBColor(255, 215, 0) 
    p.add_run('\n')
    silver = p.add_run('Silver Sponsorship\n')
    silver.bold = True
    silver.underline = True
    silver.font.color.rgb = shared.RGBColor(192, 192, 192) 
    p.add_run('\n')
    bronze = p.add_run('Bronze Sponsorship\n')
    bronze.bold = True
    bronze.underline = True
    bronze.font.color.rgb = shared.RGBColor(205, 127, 50) 
    p.add_run('\n')
    sponsorship = p.add_run('Sponsorship\n')
    sponsorship.bold = True
    sponsorship.underline = True
    p.add_run('\n')
    p.add_run('*    *    *    *\n')
    # p.add_run('Thank you very much to all who have thus far made dedications to the Shul!\n').bold = True
    # p.add_run('For available dedications, please click ').bold = True
    # p.add_hyperlink('here', address="https://www.torahtemima.org/dedications").bold = True
    # p.add_run('.\n').bold=True
    # p.add_run('*	*	*	*\n')
    # p.add_run('\n')
    # links = p.add_run('KTT Links\n')
    # links.bold = True
    # links.underline = True
    # p.add_hyperlink('KTT Whatsapp Groups', address="https://www.torahtemima.org/whatsapp").bold = True
    # p.add_run('\n')
    # p.add_hyperlink('Kiddush Sponsorship', address="https://www.torahtemima.org/form/Weekly-kiddush-sponsorships1.html").bold = True
    # p.add_run('\n')
    # p.add_hyperlink('Youth Fund', address='https://www.torahtemima.org/youth-fund').bold = True
    # p.add_run('\n')
    # p.add_hyperlink('Monthly Calendar', address='https://www.torahtemima.org/monthlycalendar').bold = True
    # p.add_run('\n')
    # p.add_hyperlink('KTT Board', address='https://www.torahtemima.org/ktt-board').bold = True
    # p.add_run('\n')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = shared.Pt(11)
    document.save(f'resources/{parsha.replace(' ','')}.docx')
